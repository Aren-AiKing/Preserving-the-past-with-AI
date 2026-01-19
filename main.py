import streamlit as st
import time
import os
from PIL import Image, ImageEnhance
from ultralytics import YOLO
import io
import zipfile
import numpy as np
import cv2
from collections import Counter
import pytesseract
from openai import OpenAI
import json
import re
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches
from skimage.measure import shannon_entropy
import traceback
from typing import Tuple, List, Optional, Dict, Any
from dataclasses import dataclass

TESSERACT_PATH = r'C:/Program Files/Tesseract-OCR/tesseract.exe'
MODEL_PATHS = {
    'detection': r"C:/Users/arene/Downloads/PageDetectNewspaper_best.pt",
    'classification': r"C:/Users/arene/Downloads/quality.pt",
    'rotation': r"C:/Users/arene/Downloads/classify_rotation.pt"
}
load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

try:
    client = OpenAI(api_key=OPENAI_API_KEY)
except Exception as e:
    st.error(f"⚠️ OpenAI initialization failed: {e}")
    client = None

@dataclass
class RotationResult:
    image: Image.Image
    was_rotated: bool
    detected_angle: str
    confidence: float
    correction_applied: int

@st.cache_resource
def load_models():
    try:
        models = {}
        for name, path in MODEL_PATHS.items():
            if not os.path.exists(path):
                st.error(f"❌ Model file not found: {path}")
                raise FileNotFoundError(f"Model not found: {path}")
            models[name] = YOLO(path)
            st.success(f"✅ Loaded {name} model")
        
        return models['detection'], models['classification'], models['rotation']
    except Exception as e:
        st.error(f"❌ Model loading failed: {str(e)}")
        raise

detection_model, classification_model, rotation_model = load_models()

def safe_wrapper(func):
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            st.error(f"❌ Error in {func.__name__}: {str(e)}")
            st.code(traceback.format_exc())
            return None
    return wrapper

@safe_wrapper
def preprocess_image(image: Image.Image, brightness: float = 1.0, 
                    contrast: float = 1.0, sharpness: float = 1.0) -> Image.Image:
    img = image.copy()
    
    if brightness != 1.0:
        enhancer = ImageEnhance.Brightness(img)
        img = enhancer.enhance(brightness)
    
    if contrast != 1.0:
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(contrast)
    
    if sharpness != 1.0:
        enhancer = ImageEnhance.Sharpness(img)
        img = enhancer.enhance(sharpness)
    
    return img

@safe_wrapper
def yolo_detect(image: Image.Image, conf_thresh: float = 0.7) -> Tuple[Image.Image, List, Any]:
    results = detection_model.predict(
        source=image, 
        conf=conf_thresh, 
        save=False, 
        imgsz=640, 
        iou=0.3, 
        verbose=False
    )
    result = results[0]
    
    has_boxes = hasattr(result, 'boxes') and result.boxes is not None and len(result.boxes) > 0
    has_masks = hasattr(result, 'masks') and result.masks is not None
    has_probs = hasattr(result, 'probs') and result.probs is not None
    
    if has_probs and not has_boxes and not has_masks:
        st.error("⚠️ Detection model is a classification model - cannot detect sections")
        return image, [], None
    
    boxes = result.boxes if has_boxes else None
    masks = result.masks.xy if (has_masks and hasattr(result.masks, 'xy')) else []
    
    result_img = result.plot()
    
    return Image.fromarray(result_img[..., ::-1]), masks, boxes

@safe_wrapper
def yolo_classify(image: Image.Image, custom_model: YOLO) -> Tuple[str, float]:
    results = custom_model.predict(source=image, save=False, verbose=False)
    result = results[0]
    
    if not (hasattr(result, 'probs') and result.probs is not None):
        st.warning("⚠️ Model returned no probabilities")
        return "unknown", 0.0
    
    class_id = int(result.probs.top1)
    confidence = float(result.probs.top1conf)
    
    if custom_model == rotation_model:
        angle = class_id * 90
        return str(angle), confidence
    
    class_name = result.names[class_id] if hasattr(result, 'names') else str(class_id)
    return class_name, confidence

def rotate_image_fixed(image: Image.Image, detected_angle: str) -> Tuple[Image.Image, int]:
    correction_map = {
        "0": 0,
        "90": -90,
        "180": 180,
        "270": 90
    }
    
    correction_angle = correction_map.get(detected_angle, 0)
    
    if correction_angle == 0:
        return image, 0
    
    rotated = image.rotate(
        correction_angle, 
        expand=True, 
        resample=Image.BICUBIC,
        fillcolor='white'
    )
    
    return rotated, correction_angle

def detect_and_rotate_image(image: Image.Image, 
                           min_confidence: float = 0.5,
                           show_debug: bool = True,
                           verify_rotation: bool = True) -> RotationResult:
    detected_angle, confidence = yolo_classify(image, rotation_model)
    
    if show_debug:
        st.info(f"🔄 Detected orientation: {detected_angle}° (confidence: {confidence:.1%})")
    
    if detected_angle == "0":
        if show_debug:
            st.success("✅ Image is correctly oriented")
        return RotationResult(
            image=image,
            was_rotated=False,
            detected_angle=detected_angle,
            confidence=confidence,
            correction_applied=0
        )
    
    if confidence < min_confidence:
        if show_debug:
            st.warning(f"⚠️ Low confidence ({confidence:.1%}) - skipping rotation")
        return RotationResult(
            image=image,
            was_rotated=False,
            detected_angle=detected_angle,
            confidence=confidence,
            correction_applied=0
        )
    
    rotated_image, correction_angle = rotate_image_fixed(image, detected_angle)
    
    if show_debug:
        st.success(f"✅ Applied {abs(correction_angle)}° {'counter-clockwise' if correction_angle > 0 else 'clockwise'} rotation")
    
    if verify_rotation:
        verify_angle, verify_conf = yolo_classify(rotated_image, rotation_model)
        
        if show_debug:
            if verify_angle == "0" and verify_conf > min_confidence:
                st.success(f"✅ Rotation verified! New orientation: {verify_angle}° ({verify_conf:.1%})")
            else:
                st.warning(f"⚠️ Post-rotation check: {verify_angle}° ({verify_conf:.1%}) - may need manual adjustment")
    
    return RotationResult(
        image=rotated_image,
        was_rotated=True,
        detected_angle=detected_angle,
        confidence=confidence,
        correction_applied=correction_angle
    )

def order_corners(pts: np.ndarray) -> np.ndarray:
    pts = np.array(pts).reshape(-1, 2).astype("float32")
    
    s = pts.sum(axis=1)
    rect = np.zeros((4, 2), dtype="float32")
    rect[0] = pts[np.argmin(s)]
    rect[2] = pts[np.argmax(s)]
    
    diff = np.diff(pts, axis=1)
    rect[1] = pts[np.argmin(diff)]
    rect[3] = pts[np.argmax(diff)]
    
    return rect

@safe_wrapper
def crop_objects(image: Image.Image, segments: List, boxes: Any = None) -> List[Tuple[str, Image.Image]]:
    crops = []
    img_np = np.array(image)

    if len(segments) == 0 and boxes is not None and len(boxes) > 0:
        st.info("📦 Using bounding boxes (no segmentation masks)")
        
        for i, box in enumerate(boxes):
            try:
                x1, y1, x2, y2 = map(int, box.xyxy[0].cpu().numpy())
                
                x1, x2 = max(0, x1), min(img_np.shape[1], x2)
                y1, y2 = max(0, y1), min(img_np.shape[0], y2)
                
                if x2 <= x1 or y2 <= y1:
                    st.warning(f"⚠️ Invalid box {i+1}: skipping")
                    continue
                
                cropped = img_np[y1:y2, x1:x2]
                pil_crop = Image.fromarray(cropped)
                crops.append((f"section_{i+1}.jpg", pil_crop))
                
            except Exception as e:
                st.warning(f"⚠️ Box {i+1} failed: {e}")
                continue
        
        return crops

    st.info("🎭 Using segmentation masks")
    
    for i, seg in enumerate(segments):
        if len(seg) < 4:
            st.warning(f"⚠️ Segment {i+1} has < 4 points: skipping")
            continue
        
        try:
            box = order_corners(seg)
            
            width = int(max(
                np.linalg.norm(box[0] - box[1]),
                np.linalg.norm(box[2] - box[3])
            ))
            height = int(max(
                np.linalg.norm(box[0] - box[3]),
                np.linalg.norm(box[1] - box[2])
            ))
            
            if width < 10 or height < 10:
                st.warning(f"⚠️ Segment {i+1} too small: {width}x{height}px")
                continue
            
            dst_pts = np.array([
                [0, 0],
                [width, 0],
                [width, height],
                [0, height]
            ], dtype="float32")
            
            M = cv2.getPerspectiveTransform(box, dst_pts)
            warped = cv2.warpPerspective(img_np, M, (width, height))
            
            pil_crop = Image.fromarray(warped)
            crops.append((f"section_{i+1}.jpg", pil_crop))
            
        except Exception as e:
            st.warning(f"⚠️ Segment {i+1} transform failed: {e}")
            continue
    
    return crops

@safe_wrapper
def ocr(image: Image.Image, lang: str = 'hye') -> str:
    return pytesseract.image_to_string(image, lang=lang)

@safe_wrapper
def get_ocr_confidence(image: Image.Image) -> float:
    data = pytesseract.image_to_data(
        image, 
        lang='hye', 
        output_type=pytesseract.Output.DICT
    )
    confidences = [conf for conf in data['conf'] if conf != -1]
    return sum(confidences) / len(confidences) if confidences else 0

@safe_wrapper
def clean_ocr_with_gpt(raw_ocr_text: str) -> str:
    if not client:
        return raw_ocr_text
    
    if not raw_ocr_text or len(raw_ocr_text.strip()) < 5:
        return raw_ocr_text
    
    response = client.chat.completions.create(
        model='gpt-4o-mini',
        messages=[
            {
                "role": "system",
                "content": "You are an Armenian language expert specializing in OCR correction. "
                          "Fix character recognition errors, correct spacing and punctuation, "
                          "maintain structure, and preserve Armenian characters. "
                          "Return only the corrected text without explanations."
            },
            {
                "role": "user",
                "content": raw_ocr_text[:4000]
            }
        ],
        temperature=0.1
    )
    return response.choices[0].message.content.strip()

@safe_wrapper
def analyze_topics(text: str) -> str:
    if not client:
        return '["unknown"]'
    
    if not text or len(text.strip()) < 10:
        return '["unknown"]'
    
    response = client.chat.completions.create(
        model='gpt-4o-mini',
        messages=[
            {
                'role': 'system',
                'content': 'You are a newspaper analysis expert for Armenian text. '
                          'Identify 1-4 main topics from: politics, sport, international, '
                          'history, science, weather, anatomy, economics, philosophy, biography. '
                          'Return ONLY a JSON array like ["politics","sport"]. '
                          'No markdown, no code blocks, just pure JSON.'
            },
            {
                'role': 'user',
                'content': text[:2000]
            }
        ],
        temperature=0
    )
    
    clean_response = response.choices[0].message.content.strip()
    clean_response = re.sub(r'^```(?:json)?\s*', '', clean_response)
    clean_response = re.sub(r'\s*```$', '', clean_response)
    
    return clean_response

@safe_wrapper
def summarize_text(text: str) -> str:
    if not client:
        return "GPT unavailable"
    
    if not text or len(text.strip()) < 10:
        return "Տեքստը չափազանց կարճ է ամփոփման համար"
    
    response = client.chat.completions.create(
        model='gpt-4o-mini',
        messages=[
            {
                'role': 'system',
                'content': 'You are an Armenian language expert. '
                          'Provide a concise summary IN ARMENIAN of the given text. '
                          'Maintain Armenian writing style and conventions.'
            },
            {
                'role': 'user',
                'content': text[:3000]
            }
        ],
        temperature=0.2
    )
    return response.choices[0].message.content.strip()

@safe_wrapper
def assess_image_quality(image: Image.Image) -> Dict[str, float]:
    quality_label, quality_prob = yolo_classify(image, classification_model)
    yolo_quality_score = quality_prob * 100
    
    img_array = np.array(image)
    gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
    
    blur_score = cv2.Laplacian(gray, cv2.CV_64F).var()
    brightness = np.mean(gray)
    contrast = np.std(gray)
    entropy = shannon_entropy(gray)
    
    blur_normalized = min(100, (blur_score / 500) * 100)
    brightness_normalized = (brightness / 255) * 100
    contrast_normalized = min(100, (contrast / 128) * 100)
    entropy_normalized = (entropy / 8) * 100
    
    overall_score = np.average([
        blur_normalized,
        brightness_normalized,
        contrast_normalized,
        entropy_normalized,
        yolo_quality_score
    ], weights=[0.3, 0.15, 0.15, 0.1, 0.3])
    
    return {
        'overall': overall_score,
        'yolo_quality': yolo_quality_score,
        'blur': blur_normalized,
        'brightness': brightness_normalized,
        'contrast': contrast_normalized,
        'entropy': entropy
    }

@safe_wrapper
def create_results_zip(uploaded_file, results_dict: Dict, docx_bytes: io.BytesIO, txt_content: str) -> io.BytesIO:
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(
            f"{uploaded_file.name}_analysis.json",
            json.dumps(results_dict, indent=2, ensure_ascii=False)
        )
        
        zf.writestr(
            f"{uploaded_file.name}_analysis.docx",
            docx_bytes.getvalue()
        )
        
        zf.writestr(
            f"{uploaded_file.name}_analysis.txt",
            txt_content
        )
    
    zip_buffer.seek(0)
    return zip_buffer

st.set_page_config(
    layout="wide",
    page_title="Newspaper Analysis Pipeline",
    page_icon="📰"
)

st.markdown("""
<style>
    .stApp { 
        max-width: 1400px; 
        margin: 0 auto; 
    }
    .main-header { 
        background: linear-gradient(135deg, #1E3D59 0%, #17A2B8 100%);
        padding: 2rem; 
        border-radius: 15px; 
        color: white; 
        margin-bottom: 2rem;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    .results-header { 
        background-color: #17A2B8; 
        color: white; 
        padding: 0.75rem 1.5rem; 
        border-radius: 8px; 
        margin-bottom: 1.5rem;
        font-weight: 600;
    }
    .metric-card {
        background-color: #f8f9fa;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #17A2B8;
    }
</style>
""", unsafe_allow_html=True)

TRANSLATIONS = {
    'English': {
        'settings': 'Settings',
        'upload_images': 'Upload Images',
        'image_preprocessing': 'Image Preprocessing',
        'brightness': 'Brightness',
        'contrast': 'Contrast',
        'sharpness': 'Sharpness',
        'detection_confidence': 'Detection Confidence',
        'rotation_confidence': 'Rotation Confidence',
        'verify_rotation': 'Verify Rotation After Correction',
        'newspaper_analysis': 'Newspaper Analysis Pipeline',
        'analysis': 'Analysis',
        'section': 'Section',
        'extracted_text': 'Extracted Text',
        'topics': 'Topics',
        'summary': 'Summary',
        'analysis_complete': '✅ Analysis complete!',
        'quality': 'Quality',
        'ocr_confidence': 'OCR Confidence'
    },
    'Armenian': {
        'settings': 'Կարգավորումներ',
        'upload_images': 'Վերբեռնել նկարներ',
        'image_preprocessing': 'Պատկերի նախնական մշակում',
        'brightness': 'Պայծառություն',
        'contrast': 'Կոնտրաստ',
        'sharpness': 'Սրություն',
        'detection_confidence': 'Հայտնաբերման վստահություն',
        'rotation_confidence': 'Պտտման վստահություն',
        'verify_rotation': 'Ստուգել պտտումը ուղղումից հետո',
        'newspaper_analysis': 'Թերթերի վերլուծության համակարգ',
        'analysis': 'Վերլուծություն',
        'section': 'Բաժին',
        'extracted_text': 'Արտահանված տեքստ',
        'topics': 'Թեմաներ',
        'summary': 'Ամփոփում',
        'analysis_complete': '✅ Վերլուծությունն ավարտված է!',
        'quality': 'Որակ',
        'ocr_confidence': 'OCR վստահություն'
    }
}

FEATURES = {
    'English': [
        "🔄 Automatic Rotation Detection & Correction",
        "📄 Page Section Detection",
        "🔤 Armenian OCR with GPT Cleaning",
        "🎯 Topic Classification",
        "📝 Automatic Summarization",
        "📊 Quality Assessment",
        "📥 Downloadable Reports (JSON, DOCX, TXT)"
    ],
    'Armenian': [
        "🔄 Ավտոմատ պտտման հայտնաբերում և ուղղում",
        "📄 Էջի բաժինների հայտնաբերում",
        "🔤 Հայերեն OCR GPT մաքրմամբ",
        "🎯 Թեմատիկ դասակարգում",
        "📝 Ավտոմատ ամփոփում",
        "📊 Որակի գնահատում",
        "📥 Ներբեռնելի հաշվետվություններ (JSON, DOCX, TXT)"
    ]
}

with st.sidebar:
    st.markdown('<div class="main-header"><h2>⚙️ Settings</h2></div>', unsafe_allow_html=True)
    
    language = st.selectbox("Language / Լեզու", ["English", "Armenian"])
    t = TRANSLATIONS[language]
    
    st.markdown("---")
    
    uploaded_files = st.file_uploader(
        t['upload_images'],
        type=["jpg", "png", "jpeg"],
        accept_multiple_files=True,
        help="Upload one or more newspaper images"
    )
    
    st.markdown("---")
    
    st.subheader(t['image_preprocessing'])
    brightness = st.slider(t['brightness'], 0.0, 2.0, 1.0, 0.1)
    contrast = st.slider(t['contrast'], 0.0, 2.0, 1.0, 0.1)
    sharpness = st.slider(t['sharpness'], 0.0, 2.0, 1.0, 0.1)
    
    st.markdown("---")
    
    st.subheader("🔍 Detection Settings")
    confidence = st.slider(t['detection_confidence'], 0.1, 1.0, 0.3, 0.05)
    
    st.markdown("---")
    
    st.subheader("🔄 Rotation Settings")
    rotation_confidence = st.slider(t['rotation_confidence'], 0.1, 1.0, 0.5, 0.05)
    verify_rotation = st.checkbox(t['verify_rotation'], value=True)

st.markdown(f'<div class="main-header"><h1>📰 {t["newspaper_analysis"]}</h1></div>', unsafe_allow_html=True)

if not uploaded_files:
    st.info("👆 Please upload newspaper images using the sidebar")
    st.markdown("### Features / Հնարավորություններ")
    
    for feature in FEATURES[language]:
        st.markdown(f"- {feature}")

else:
    for uploaded_file in uploaded_files:
        st.markdown("---")
        st.markdown(f"## 📄 {uploaded_file.name}")
        
        try:
            selected_image = Image.open(uploaded_file).convert("RGB")
            
            with st.spinner("🔧 Preprocessing image..."):
                processed_image = preprocess_image(selected_image, brightness, contrast, sharpness)
            
            col1, col2 = st.columns([1, 1])
            
            with col1:
                st.markdown("### Original / Preprocessed")
                st.image(processed_image, use_container_width=True)
            
            with col2:
                if st.button(f"🔬 Start Analysis", key=f"btn_{uploaded_file.name}", type="primary", use_container_width=True):
                    
                    with st.spinner("📊 Checking quality..."):
                        try:
                            quality_label, quality_prob = yolo_classify(processed_image, classification_model)
                            st.info(f"📊 **Quality Classification:** {quality_label} ({quality_prob:.1%})")
                        except Exception as e:
                            st.warning(f"Quality check failed: {e}")
                    
                    with st.spinner("🔄 Detecting and correcting rotation..."):
                        try:
                            rotation_result = detect_and_rotate_image(
                                processed_image,
                                min_confidence=rotation_confidence,
                                show_debug=True,
                                verify_rotation=verify_rotation
                            )
                            
                            processed_image = rotation_result.image
                            
                            if rotation_result.was_rotated:
                                st.image(
                                    processed_image,
                                    caption=f"After {rotation_result.correction_applied}° rotation",
                                    use_container_width=True
                                )
                        except Exception as e:
                            st.warning(f"Rotation check failed: {e}")
                            st.code(traceback.format_exc())
                    
                    with st.spinner("🔍 Detecting sections..."):
                        try:
                            det_img, masks, boxes = yolo_detect(processed_image, confidence)
                            
                            if det_img:
                                st.image(det_img, caption="Detection Results", use_container_width=True)
                        except Exception as e:
                            st.error(f"Detection failed: {e}")
                            st.code(traceback.format_exc())
                            continue
                    
                    with st.spinner("✂️ Extracting sections..."):
                        try:
                            crops = crop_objects(processed_image, masks, boxes)
                        except Exception as e:
                            st.error(f"Cropping failed: {e}")
                            crops = []
                    
                    if not crops:
                        st.error(f"⚠️ No sections detected (confidence={confidence:.2f})")
                        st.info("💡 Try lowering the detection confidence slider to 0.1-0.2")
                        continue
                    
                    st.success(f"✅ Found {len(crops)} sections")
            
            if 'crops' in locals() and crops:
                st.markdown(f'<div class="results-header"><h2>📋 {t["analysis"]} Results</h2></div>', unsafe_allow_html=True)
                
                all_sections_data = []
                
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                for idx, (name, crop) in enumerate(crops, 1):
                    progress_bar.progress(idx / len(crops))
                    status_text.text(f"Processing section {idx}/{len(crops)}...")
                    
                    try:
                        with st.expander(f"📄 {t['section']} {idx}", expanded=(idx == 1)):
                            col_img, col_text = st.columns([1, 2])
                            
                            with col_img:
                                st.image(crop, caption=f"{t['section']} {idx}", use_container_width=True)
                                
                                quality_scores = assess_image_quality(crop)
                                
                                st.markdown("#### Quality Metrics")
                                st.metric("Overall Quality", f"{quality_scores['overall']:.0f}%")
                                
                                with st.expander("📊 Detailed Metrics"):
                                    st.metric("YOLO Quality", f"{quality_scores['yolo_quality']:.0f}%")
                                    st.metric("Sharpness", f"{quality_scores['blur']:.0f}%")
                                    st.metric("Brightness", f"{quality_scores['brightness']:.0f}%")
                                    st.metric("Contrast", f"{quality_scores['contrast']:.0f}%")
                            
                            with col_text:
                                with st.spinner("🔤 Extracting text..."):
                                    raw_text = ocr(crop)
                                    ocr_conf = get_ocr_confidence(crop)
                                    cleaned_text = clean_ocr_with_gpt(raw_text)
                                
                                tab1, tab2 = st.tabs([t['extracted_text'], t['analysis']])
                                
                                with tab1:
                                    st.progress(ocr_conf / 100, text=f"{t['ocr_confidence']}: {ocr_conf:.0f}%")
                                    st.text_area(
                                        'Extracted Text',
                                        value=cleaned_text,
                                        height=300,
                                        key=f"txt_{idx}_{uploaded_file.name}"
                                    )
                                
                                with tab2:
                                    with st.spinner("🎯 Analyzing topics..."):
                                        themes = analyze_topics(cleaned_text)
                                    
                                    st.markdown(f"**{t['topics']}:**")
                                    try:
                                        topics_list = json.loads(themes)
                                        for topic in topics_list:
                                            st.badge(topic)
                                    except:
                                        st.text(themes)
                                    
                                    with st.spinner("📝 Generating summary..."):
                                        summary = summarize_text(cleaned_text)
                                    
                                    st.markdown(f"**{t['summary']}:**")
                                    st.text_area(
                                        "Summary",
                                        value=summary,
                                        height=200,
                                        key=f"sum_{idx}_{uploaded_file.name}"
                                    )
                                
                                all_sections_data.append({
                                    "section": idx,
                                    "text": cleaned_text,
                                    "themes": themes,
                                    "summary": summary,
                                    "quality_scores": quality_scores,
                                    "ocr_confidence": ocr_conf
                                })
                    
                    except Exception as e:
                        st.error(f"❌ Section {idx} processing failed: {e}")
                        st.code(traceback.format_exc())
                        continue
                
                progress_bar.empty()
                status_text.empty()
                
                if all_sections_data:
                    try:
                        with st.spinner("📦 Creating download package..."):
                            results_dict = {
                                "filename": uploaded_file.name,
                                "rotation_applied": rotation_result.was_rotated if 'rotation_result' in locals() else False,
                                "rotation_angle": rotation_result.correction_applied if 'rotation_result' in locals() else 0,
                                "sections": all_sections_data
                            }
                            
                            doc = Document()
                            doc.add_heading(f'Analysis Results: {uploaded_file.name}', 0)
                            
                            if 'rotation_result' in locals() and rotation_result.was_rotated:
                                doc.add_paragraph(f'Rotation applied: {rotation_result.correction_applied}°')
                            
                            for sd in all_sections_data:
                                doc.add_heading(f'Section {sd["section"]}', 1)
                                
                                doc.add_heading('Quality Metrics', 2)
                                doc.add_paragraph(f'Overall: {sd["quality_scores"]["overall"]:.1f}%')
                                doc.add_paragraph(f'OCR Confidence: {sd["ocr_confidence"]:.1f}%')
                                
                                doc.add_heading('Extracted Text', 2)
                                doc.add_paragraph(sd["text"])
                                
                                doc.add_heading('Topics', 2)
                                try:
                                    doc.add_paragraph(', '.join(json.loads(sd["themes"])))
                                except:
                                    doc.add_paragraph(sd["themes"])
                                
                                doc.add_heading('Summary', 2)
                                doc.add_paragraph(sd["summary"])
                                
                                doc.add_page_break()
                            
                            docx_bytes = io.BytesIO()
                            doc.save(docx_bytes)
                            docx_bytes.seek(0)
                            
                            txt_content = f"Analysis Results: {uploaded_file.name}\n{'='*60}\n\n"
                            
                            if 'rotation_result' in locals() and rotation_result.was_rotated:
                                txt_content += f"Rotation applied: {rotation_result.correction_applied}°\n\n"
                            
                            for sd in all_sections_data:
                                txt_content += f"\n{'='*60}\n"
                                txt_content += f"Section {sd['section']}\n"
                                txt_content += f"{'='*60}\n\n"
                                txt_content += f"Quality: {sd['quality_scores']['overall']:.1f}%\n"
                                txt_content += f"OCR Confidence: {sd['ocr_confidence']:.1f}%\n\n"
                                txt_content += f"Text:\n{sd['text']}\n\n"
                                txt_content += f"Topics:\n{sd['themes']}\n\n"
                                txt_content += f"Summary:\n{sd['summary']}\n\n"
                            
                            results_zip = create_results_zip(uploaded_file, results_dict, docx_bytes, txt_content)
                            
                            if results_zip:
                                st.download_button(
                                    "📥 Download Complete Analysis",
                                    data=results_zip.getvalue(),
                                    file_name=f"{uploaded_file.name}_analysis.zip",
                                    mime="application/zip",
                                    use_container_width=True,
                                    type="primary"
                                )
                            
                            st.success(t['analysis_complete'])
                    
                    except Exception as e:
                        st.error(f"❌ Download creation failed: {e}")
                        st.code(traceback.format_exc())
        
        except Exception as e:
            st.error(f"❌ Fatal error processing {uploaded_file.name}")
            st.error(str(e))
            with st.expander("🐛 Debug Information"):
                st.code(traceback.format_exc())
