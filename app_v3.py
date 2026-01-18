import streamlit as st
import time
import os
from dotenv import load_dotenv
from PIL import Image, ImageEnhance
from ultralytics import YOLO
import io
import zipfile
import numpy as np
import cv2
from collections import Counter
import pytesseract
from openai import OpenAI
import json
import re
from docx import Document
from docx.shared import Inches
from skimage.measure import shannon_entropy

pytesseract.pytesseract.tesseract_cmd = r'C:/Program Files/Tesseract-OCR/tesseract.exe'
load_dotenv()
client = OpenAI(os.getenv("api_key"))

@st.cache_data
def load_models():
    detection_model = YOLO("C:/Users/arene/Downloads/PageDetectNewspaper_best.pt")
    classification_model = YOLO("C:/Users/arene/Downloads/quality.pt")
    rotation_model = YOLO("C:/Users/arene/Downloads/classify_rotation.pt")
    return detection_model, classification_model, rotation_model

detection_model, classification_model, rotation_model = load_models()

def preprocess_image(image, brightness=1.0, contrast=1.0, sharpness=1.0):
    img = image
    if brightness != 1.0:
        img = ImageEnhance.Brightness(img).enhance(brightness)
    if contrast != 1.0:
        img = ImageEnhance.Contrast(img).enhance(contrast)
    if sharpness != 1.0:
        img = ImageEnhance.Sharpness(img).enhance(sharpness)
    return img

def yolo_detect(image, conf_thresh=0.7):
    results = detection_model.predict(source=image, conf=conf_thresh, save=False, imgsz=640, iou=0.3)
    result_img = results[0].plot()
    masks = results[0].masks.xy if results[0].masks else []
    return Image.fromarray(result_img[..., ::-1]), masks

def yolo_classify(image, custom_model):
    results = custom_model.predict(source=image, save=False)[0]
    class_id = int(results.probs.top1)
    class_name = results.names[class_id] if custom_model != rotation_model else str(class_id * 90)
    prob = float(results.probs.top1conf)
    return class_name, prob

def rotate_image(image, orientation_label):
    rotation_angles = {
        "0": 0,     # No rotation
        "90": 90,   # Rotate 90 degrees clockwise
        "180": 180, # Rotate 180 degrees
        "270": -90  # Rotate 270 degrees clockwise (same as -90)
    }
    angle = rotation_angles.get(orientation_label, 0)
    if angle != 0:
        return image.rotate(angle, expand=True)
    return image

def order_corners(pts):
    pts = np.array(pts).reshape(-1, 2).astype("float32")
    s = pts.sum(axis=1)
    diff = np.diff(pts, axis=1)
    rect = np.zeros((4, 2), dtype="float32")
    rect[0] = pts[np.argmin(s)]
    rect[1] = pts[np.argmin(diff)]
    rect[2] = pts[np.argmax(s)]
    rect[3] = pts[np.argmax(diff)]
    return rect

def crop_objects(image, segments):
    crops = []
    img_np = np.array(image)

    for i, seg in enumerate(segments):
        if len(seg) < 4:
            continue
        try:
            box = order_corners(seg)
            width = int(max(np.linalg.norm(box[0] - box[1]), np.linalg.norm(box[2] - box[3])))
            height = int(max(np.linalg.norm(box[0] - box[3]), np.linalg.norm(box[1] - box[2])))
            dst_pts = np.array([[0, 0], [width, 0], [width, height], [0, height]], dtype="float32")
            M = cv2.getPerspectiveTransform(box, dst_pts)
            warped = cv2.warpPerspective(img_np, M, (width, height))
            pil_crop = Image.fromarray(warped)
            crops.append((f"object_{i+1}.jpg", pil_crop))
        except Exception as e:
            print(f"Polygon {i} failed: {e}")
            continue
    return crops

def create_results_zip(uploaded_file, results_dict, docx_bytes, txt_content):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(
            f"{uploaded_file.name}_analysis.json", 
            json.dumps(results_dict, indent=2)
        )
        
        zf.writestr(
            f"{uploaded_file.name}_analysis.docx",
            docx_bytes.getvalue()
        )
        
        zf.writestr(
            f"{uploaded_file.name}_analysis.txt",
            txt_content
        )
    
    zip_buffer.seek(0)
    return zip_buffer

def count_classes(class_names):
    return Counter(class_names)

def ocr(image,lang='hye'):
    text = pytesseract.image_to_string(image,lang=lang)
    return text

def get_ocr_confidence(image):
    data = pytesseract.image_to_data(image, lang='hye', output_type=pytesseract.Output.DICT)
    confidences = [conf for conf in data['conf'] if conf != -1]
    return sum(confidences) / len(confidences) if confidences else 0

def analyze_topics(text):
    message = [
        {
            'role': 'system',
            'content': (
                'You are a newspapers analysis expert analyzing Armenian text. '
                'Based on the Armenian text input, identify the main topics. '
                'Return ONLY a valid JSON array with 1-4 topics from: '
                '["politics","sport","international","history","science","weather",'
                '"anatomy","economics","philosophy","biography"]. '
                'No formatting, no backticks, no markdown - just the pure JSON array.'
            )
        },
        {'role': 'user', 'content': text}
    ]

    response = client.chat.completions.create(
        model='gpt-4o-mini-2024-07-18',
        messages=message,
        temperature=0
    )
    
    clean_response = response.choices[0].message.content.strip()
    clean_response = re.sub(r'^```.*\n?', '', clean_response)
    clean_response = re.sub(r'\n```$', '', clean_response)
    return clean_response

def clean_ocr_with_gpt(raw_ocr_text):
    message = [
        {
            "role": "system",
            "content": (
                "You are an Armenian language expert specializing in OCR correction. "
                "Given raw OCR output from a scanned Armenian newspaper, your task is to:\n"
                "1. Fix character recognition errors\n"
                "2. Correct word spacing and punctuation\n"
                "3. Maintain sentence structure and formatting\n"
                "4. Preserve Armenian diacritics and special characters\n"
                "Only return the corrected text without any explanations."
            )
        },
        {"role": "user", "content": raw_ocr_text}
    ]

    response = client.chat.completions.create(
        model='gpt-4o-mini-2024-07-18',
        messages=message,
        temperature=0.1
    )

    return response.choices[0].message.content.strip()

def summarize_text(text):
    messages = [
        {
            'role': 'system',
            'content': (
                'You are an Armenian language expert that summarizes content. '
                'Given Armenian text, provide a concise summary IN ARMENIAN. '
                'Keep the summary in Armenian language and maintain Armenian '
                'writing style and conventions.'
            )
        },
        {
            'role': 'user',
            'content': text
        }
    ]
    
    response = client.chat.completions.create(
        model = 'gpt-4o-mini-2024-07-18',
        messages = messages,
        temperature = 0.2
    )

    return response.choices[0].message.content.strip()

def assess_image_quality(image):
    quality_label, quality_prob = yolo_classify(image, classification_model)
    yolo_quality_score = quality_prob * 100
    
    img_array = np.array(image)
    gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
    
    blur_score = cv2.Laplacian(gray, cv2.CV_64F).var()
    brightness = np.mean(gray)
    contrast = np.std(gray)
    entropy = shannon_entropy(gray)
    
    blur_score = min(100, (blur_score / 500) * 100)
    brightness_score = (brightness / 255) * 100
    contrast_score = min(100, (contrast / 128) * 100)
    entropy_score = (entropy / 8) * 100
    
    overall_score = np.mean([
        blur_score * 0.3,
        brightness_score * 0.15,
        contrast_score * 0.15,
        entropy_score * 0.1,
        yolo_quality_score * 0.3
    ])
    
    return {
        'yolo_quality': yolo_quality_score,
        'blur': blur_score,
        'brightness': brightness_score,
        'contrast': contrast_score,
        'entropy': entropy,
        'overall': overall_score
    }

st.set_page_config(layout="wide", page_title="Advanced Newspaper Analysis")
st.markdown("""
<style>
    .stApp {
        max-width: 1200px;
        margin: 0 auto;
    }
    .main-header {
        background-color: #1E3D59;
        padding: 1.5rem;
        border-radius: 10px;
        color: white;
        margin-bottom: 2rem;
    }
    .results-header {
        background-color: #17A2B8;
        color: white;
        padding: 0.5rem 1rem;
        border-radius: 5px;
        margin-bottom: 1rem;
    }
    .stTextArea textarea {
        font-size: 16px !important;
        line-height: 1.5 !important;
    }
</style>
""", unsafe_allow_html=True)

TRANSLATIONS = {
    'English': {
        'settings': 'Settings',
        'upload_images': 'Upload Images',
        'image_preprocessing': 'Image Preprocessing',
        'brightness': 'Brightness',
        'contrast': 'Contrast',
        'sharpness': 'Sharpness',
        'detection_confidence': 'Detection Confidence',
        'newspaper_analysis': 'Newspaper Analysis Pipeline',
        'pipeline_steps': 'Upload → Preprocess → Classify → Detect → Analyze',
        'processing_image': 'Processing image...',
        'running_analysis': 'Running analysis...',
        'classifying': 'Classifying image...',
        'detecting_orientation': 'Detecting orientation...',
        'detecting_sections': 'Detecting sections...',
        'processing_sections': 'Processing detected sections...',
        'analysis_results': 'Analysis Results',
        'section': 'Section',
        'extracted_text': 'Extracted Text',
        'analysis': 'Analysis',
        'topics': 'Topics',
        'summary': 'Summary',
        'analysis_complete': '✅ Analysis complete!',
        'download_results': '📥 Download Analysis Results',
    },
    'Armenian': {
        'settings': 'Կարգավորումներ',
        'upload_images': 'Վերբեռնել նկարներ',
        'image_preprocessing': 'Պատկերի նախնական մշակում',
        'brightness': 'Պայծառություն',
        'contrast': 'Կոնտրաստ',
        'sharpness': 'Սրություն',
        'detection_confidence': 'Հայտնաբերման վստահություն',
        'newspaper_analysis': 'Թերթերի վերլուծության համակարգ',
        'pipeline_steps': 'Վերբեռնել → Մշակել → Դասակարգել → Հայտնաբերել → Վերլուծել',
        'processing_image': 'Նկարի մշակում...',
        'running_analysis': 'Վերլուծություն...',
        'classifying': 'Դասակարգում...',
        'detecting_orientation': 'Կողմնորոշման հայտնաբերում...',
        'detecting_sections': 'Բաժինների հայտնաբերում...',
        'processing_sections': 'Բաժինների մշակում...',
        'analysis_results': 'Վերլուծության արդյունքներ',
        'section': 'Բաժին',
        'extracted_text': 'Արտահանված տեքստ',
        'analysis': 'Վերլուծություն',
        'topics': 'Թեմաներ',
        'summary': 'Ամփոփում',
        'analysis_complete': '✅ Վերլուծությունն ավարտված է!',
        'download_results': '📥 Ներբեռնել արդյունքները',
    }
}

st.markdown("""
<style>
    .stProgress > div > div > div > div {
        background-image: linear-gradient(to right, #ff4b4b, #f0b907, #48c78e);
    }
    .confidence-label {
        color: #666;
        font-size: 14px;
        margin-bottom: 5px;
    }
    .confidence-container {
        background: #f0f2f6;
        padding: 10px;
        border-radius: 5px;
        margin: 10px 0;
    }
</style>
""", unsafe_allow_html=True)

with st.sidebar:
    st.markdown('<div class="main-header">Settings</div>', unsafe_allow_html=True)
    language = st.selectbox("Language / Լեզու", ["English", "Armenian"])
    t = TRANSLATIONS[language]
    
    uploaded_files = st.file_uploader(t['upload_images'], type=["jpg", "png", "jpeg"], accept_multiple_files=True)
    
    st.markdown('<div class="section-box">', unsafe_allow_html=True)
    st.subheader(t['image_preprocessing'])
    brightness = st.slider(t['brightness'], 0.0, 2.0, 1.0)
    contrast = st.slider(t['contrast'], 0.0, 2.0, 1.0)
    sharpness = st.slider(t['sharpness'], 0.0, 2.0, 1.0)
    confidence = st.slider(t['detection_confidence'], 0.1, 1.0, 0.7)
    st.markdown('</div>', unsafe_allow_html=True)

st.markdown(f'<div class="main-header"><h1>{t["newspaper_analysis"]}</h1></div>', unsafe_allow_html=True)
st.markdown(f"### {t['pipeline_steps']}")

if uploaded_files:
    for uploaded_file in uploaded_files:
        with st.spinner(t['processing_image']):
            try:
                selected_image = Image.open(uploaded_file).convert("RGB")
                processed_image = preprocess_image(selected_image, brightness, contrast, sharpness)
                
                col1, col2 = st.columns(2)
                with col1:
                    st.image(processed_image, caption="Processed Image")
                
                if st.button(f"{t['analysis']} {uploaded_file.name}", key=f"analyze_{uploaded_file.name}"):
                    with st.status(t['running_analysis']) as status:
                        status.write(t['classifying'])
                        label, prob = yolo_classify(processed_image, classification_model)
                        
                        status.write(t['detecting_orientation'])
                        orientation_label, orientation_prob = yolo_classify(processed_image, rotation_model)
                        
                        if orientation_prob > 0.5:  # Only rotate if confidence is high enough
                            processed_image = rotate_image(processed_image, orientation_label)
                            st.info(f"Image rotated {orientation_label}° (confidence: {orientation_prob:.2%})")
                        
                        status.write(t['detecting_sections'])
                        det_img, masks = yolo_detect(processed_image, confidence)
                        
                        with col2:
                            st.image(det_img, caption="Detection Results")
                        
                        status.write(t['processing_sections'])
                        crops = crop_objects(processed_image, masks)
                        
                        if crops:
                            st.markdown(f'<div class="results-header"><h2>{t["analysis_results"]}</h2></div>', unsafe_allow_html=True)
                            for idx, (name, crop) in enumerate(crops, 1):
                                st.markdown(f'<div class="section-box">', unsafe_allow_html=True)
                                st.markdown(f"#### {t['section']} {idx}")
                                col_img, col_text = st.columns([1, 2])
                                
                                with col_img:
                                    st.image(crop, caption=f"Section {idx}")
                                    
                                    quality_scores = assess_image_quality(crop)
                                    
                                    st.markdown("### Image Quality Metrics")
                                    q_col1, q_col2 = st.columns(2)
                                    
                                    with q_col1:
                                        st.metric("Overall Quality", f"{quality_scores['overall']:.1f}%")
                                        st.metric("YOLO Quality", f"{quality_scores['yolo_quality']:.1f}%")
                                        st.metric("Blur Score", f"{quality_scores['blur']:.1f}%")
                                    
                                    with q_col2:
                                        st.metric("Brightness", f"{quality_scores['brightness']:.1f}%")
                                        st.metric("Contrast", f"{quality_scores['contrast']:.1f}%")
                                        st.metric("Detail (Entropy)", f"{quality_scores['entropy']:.1f}%")
                                    
                                    if quality_scores['overall'] < 50:
                                        st.warning("⚠️ Low image quality detected. Results may be affected.")
                                    elif quality_scores['overall'] < 70:
                                        st.info("ℹ️ Moderate image quality. Some results may be affected.")
                                    else:
                                        st.success("✅ Good image quality")
                                
                                with col_text:
                                    with st.spinner("Analyzing content..."):
                                        raw_text = ocr(crop)
                                        cleaned_text = clean_ocr_with_gpt(raw_text)
                                        ocr_confidence = get_ocr_confidence(crop)
                                        
                                        text_tab, analysis_tab = st.tabs([t['extracted_text'], t['analysis']])
                                        
                                        with text_tab:
                                            st.markdown('<div class="confidence-container">', unsafe_allow_html=True)
                                            st.markdown('<p class="confidence-label">OCR Confidence</p>', unsafe_allow_html=True)
                                            st.progress(ocr_confidence / 100)
                                            st.text_area(
                                                'Cleaned Text', 
                                                value=cleaned_text, 
                                                height=400,
                                                key=f"text_{idx}"
                                            )
                                            st.markdown('</div>', unsafe_allow_html=True)
                                        
                                        with analysis_tab:
                                            theme = analyze_topics(cleaned_text)
                                            summary = summarize_text(cleaned_text)
                                            
                                            summary_confidence = min(len(summary) / len(cleaned_text) * 100, 95)
                                            
                                            st.markdown('<div class="confidence-container">', unsafe_allow_html=True)
                                            st.markdown('<p class="confidence-label">Summary Confidence</p>', unsafe_allow_html=True)
                                            st.progress(summary_confidence / 100)
                                            
                                            st.markdown(f"**{t['topics']}:**")
                                            try:
                                                parsed_theme = json.loads(theme)
                                                st.json(parsed_theme)
                                            except json.JSONDecodeError as e:
                                                st.error(f"Error parsing topics: {str(e)}")
                                            
                                            st.markdown(f"**{t['summary']}:**")
                                            st.text_area(
                                                "Ամփոփում",
                                                value=summary,
                                                height=250,
                                                key=f"summary_{idx}"
                                            )
                                            st.markdown('</div>', unsafe_allow_html=True)
                                st.markdown('</div>', unsafe_allow_html=True)
                            
                            st.markdown("""
                            <style>
                                .stDownloadButton {
                                    background-color: #28A745 !important;
                                    color: white !important;
                                    padding: 0.5rem 1rem !important;
                                    border-radius: 5px !important;
                                    margin-top: 2rem !important;
                                }
                            </style>
                            """, unsafe_allow_html=True)
                            
                            results_dict = {
                                "filename": uploaded_file.name,
                                "sections": [{
                                    "section": idx,
                                    "text": cleaned_text,
                                    "themes": theme,
                                    "summary": summary,
                                    "quality_scores": quality_scores
                                } for idx, ((_, _), cleaned_text, theme, summary, quality_scores) in 
                                    enumerate(zip(crops, [cleaned_text]*len(crops), 
                                                [theme]*len(crops), [summary]*len(crops),
                                                [quality_scores]*len(crops)), 1)]
                            }
                            
                            doc = Document()
                            doc.add_heading(f'Analysis Results for {uploaded_file.name}', 0)
                            
                            for idx, (cleaned_text, theme, summary) in enumerate(
                                zip([cleaned_text]*len(crops), 
                                    [theme]*len(crops), [summary]*len(crops)), 1):
                                
                                doc.add_heading(f'Section {idx}', level=1)
                                doc.add_heading('Extracted Text', level=2)
                                doc.add_paragraph(cleaned_text)
                                
                                doc.add_heading('Topics', level=2)
                                try:
                                    parsed_theme = json.loads(theme)
                                    doc.add_paragraph(', '.join(parsed_theme))
                                except:
                                    doc.add_paragraph(theme)
                                
                                doc.add_heading('Summary', level=2)
                                doc.add_paragraph(summary)
                                doc.add_page_break()

                            docx_bytes = io.BytesIO()
                            doc.save(docx_bytes)
                            docx_bytes.seek(0)
                            
                            txt_content = f"Analysis Results for {uploaded_file.name}\n\n"
                            for idx, ((_, crop), cleaned_text, theme, summary) in enumerate(
                                zip(crops, [cleaned_text]*len(crops), 
                                    [theme]*len(crops), [summary]*len(crops)), 1):
                                txt_content += f"\nSection {idx}\n"
                                txt_content += f"{'='*50}\n"
                                txt_content += f"Extracted Text:\n{cleaned_text}\n\n"
                                txt_content += f"Topics:\n{theme}\n\n"
                                txt_content += f"Summary:\n{summary}\n\n"
                            
                            results_zip = create_results_zip(
                                uploaded_file,
                                results_dict,
                                docx_bytes,
                                txt_content
                            )
                            
                            download_container = st.container()
                            with download_container:
                                st.download_button(
                                    "📥 Download Complete Analysis",
                                    data=results_zip.getvalue(),
                                    file_name=f"{uploaded_file.name}_analysis.zip",
                                    mime="application/zip",
                                    use_container_width=True,
                                    help="Download a ZIP file containing the complete analysis including text, summaries, and metrics"
                                )

                        status.update(label=t['analysis_complete'], state="complete")
            
            except Exception as e:
                st.error(f"Error processing {uploaded_file.name}: {str(e)}")
                continue
